"""
OCRVerifier v1 — Offline Document Analysis
==========================================
PIPELINE:
  1. Preprocess image with OpenCV (deskew, denoise, threshold)
  2. Extract raw text with pytesseract
  3. Parse: ID type, DOB, Name, Vehicle No, Diagnosis etc.
  4. Compare DOB with stated age
  5. Return structured result — no paid API needed

SUPPORTED DOCUMENTS:
  - Aadhaar Card
  - PAN Card
  - Driving Licence
  - Passport
  - Voter ID
  - Health Report
  - RC Book (Vehicle)
  - Previous Insurance Policy
  - PDF (any of above)
  - DOCX / TXT

After offline verification → Gemini AI takes over conversation
"""

import os, re, io, datetime, logging, json
import numpy as np

log = logging.getLogger("OCRVerifier")

# ── Safe imports (graceful fallback if package missing) ───────────────────────
try:
    import cv2
    CV2_OK = True
except ImportError:
    CV2_OK = False
    log.warning("[OCR] OpenCV not available — using PIL only")

try:
    import pytesseract
    from pytesseract import Output
    TESS_OK = True
    # Try to locate tesseract binary
    for _p in ["/usr/bin/tesseract", "/usr/local/bin/tesseract",
               r"C:\Program Files\Tesseract-OCR\tesseract.exe",
               r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"]:
        if os.path.exists(_p):
            pytesseract.pytesseract.tesseract_cmd = _p
            break
except ImportError:
    TESS_OK = False
    log.warning("[OCR] pytesseract not available")

try:
    from PIL import Image, ImageEnhance, ImageFilter
    PIL_OK = True
except ImportError:
    PIL_OK = False

try:
    import pypdf as PyPDF2
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    import docx
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


# ─────────────────────────────────────────────────────────────────────────────
#  IMAGE PREPROCESSOR (OpenCV)
# ─────────────────────────────────────────────────────────────────────────────
class ImagePreprocessor:
    """Applies OCR-optimized preprocessing using OpenCV + PIL."""

    @staticmethod
    def load_bytes(file_bytes: bytes) -> "np.ndarray | None":
        if CV2_OK:
            arr = np.frombuffer(file_bytes, np.uint8)
            img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
            return img
        elif PIL_OK:
            pil = Image.open(io.BytesIO(file_bytes)).convert("RGB")
            return np.array(pil)
        return None

    @staticmethod
    def preprocess(img: "np.ndarray") -> "list[np.ndarray]":
        """Returns multiple versions of the image for best OCR coverage."""
        if img is None:
            return []
        versions = []

        if CV2_OK:
            # Version 1: Grayscale + Otsu threshold (best for printed text)
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            _, thresh1 = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            versions.append(thresh1)

            # Version 2: Adaptive threshold (good for uneven lighting)
            thresh2 = cv2.adaptiveThreshold(
                gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY, 11, 2
            )
            versions.append(thresh2)

            # Version 3: Denoised grayscale (good for noisy images)
            denoised = cv2.fastNlMeansDenoising(gray, h=10)
            versions.append(denoised)

            # Version 4: Sharpened (helps with blurry text)
            kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]])
            sharpened = cv2.filter2D(gray, -1, kernel)
            versions.append(sharpened)

            # Version 5: Upscaled (helps small text like DOB on Aadhaar)
            h, w = gray.shape
            scaled = cv2.resize(gray, (w * 2, h * 2), interpolation=cv2.INTER_CUBIC)
            versions.append(scaled)

        elif PIL_OK:
            # Fallback: PIL enhancement
            pil = Image.fromarray(img).convert("L")
            versions.append(np.array(pil))
            enhanced = ImageEnhance.Contrast(pil).enhance(2.0)
            versions.append(np.array(enhanced))

        return versions

    @staticmethod
    def assess_quality(img: "np.ndarray") -> dict:
        """Quick quality check before OCR."""
        if img is None or not CV2_OK:
            return {"quality": "unknown", "score": 50}

        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY) if len(img.shape) == 3 else img

        # Blurriness (Laplacian variance)
        blur_score = cv2.Laplacian(gray, cv2.CV_64F).var()

        # Brightness
        brightness = np.mean(gray)

        # Determine quality
        if blur_score < 50:
            quality = "blurry"
        elif brightness < 40:
            quality = "dark"
        elif brightness > 240:
            quality = "overexposed"
        elif blur_score > 100 and 60 < brightness < 220:
            quality = "good"
        else:
            quality = "acceptable"

        return {
            "quality": quality,
            "blur_score": round(blur_score, 1),
            "brightness": round(float(brightness), 1),
            "score": min(100, int(blur_score / 2 + (brightness / 2.55) * 0.3))
        }


# ─────────────────────────────────────────────────────────────────────────────
#  TEXT EXTRACTOR
# ─────────────────────────────────────────────────────────────────────────────
class TextExtractor:
    """Extracts raw text from images, PDFs, DOCX, TXT using pytesseract."""

    # Tesseract configs optimized for different document types
    CONFIGS = [
        "--oem 3 --psm 3",   # Auto page segmentation (default)
        "--oem 3 --psm 6",   # Uniform block of text
        "--oem 3 --psm 4",   # Single column
        "--oem 1 --psm 3",   # LSTM only
    ]

    @classmethod
    def from_image_bytes(cls, file_bytes: bytes, ext: str) -> str:
        """Extract text from image file bytes."""
        if not TESS_OK:
            return ""

        img = ImagePreprocessor.load_bytes(file_bytes)
        if img is None:
            return ""

        versions = ImagePreprocessor.preprocess(img)
        best_text = ""
        best_len  = 0

        for version in versions:
            for config in cls.CONFIGS:
                try:
                    if PIL_OK:
                        pil_img = Image.fromarray(version)
                        text = pytesseract.image_to_string(pil_img, lang="eng", config=config)
                    else:
                        text = pytesseract.image_to_string(version, lang="eng", config=config)

                    text = text.strip()
                    # Pick the result with most usable characters
                    usable = sum(1 for c in text if c.isalnum() or c in "/-:")
                    if usable > best_len:
                        best_len  = usable
                        best_text = text
                except Exception as e:
                    log.debug(f"[OCR] Config {config} failed: {e}")
                    continue

        log.info(f"[OCR] Extracted {len(best_text)} chars from image")
        return best_text

    @classmethod
    def from_pdf_bytes(cls, file_bytes: bytes) -> str:
        """Extract text from PDF — tries text layer first, then image OCR."""
        texts = []

        # Method 1: Extract text layer (fast, works for digital PDFs)
        if PDF_OK:
            try:
                reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                for page in reader.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        texts.append(t.strip())
                if texts:
                    combined = "\n".join(texts)
                    log.info(f"[OCR] PDF text layer: {len(combined)} chars")
                    return combined
            except Exception as e:
                log.debug(f"[OCR] PDF text extract failed: {e}")

        # Method 2: Convert PDF pages to images then OCR
        try:
            from pdf2image import convert_from_bytes
            pages = convert_from_bytes(file_bytes, dpi=200)
            for page_img in pages:
                img_bytes = io.BytesIO()
                page_img.save(img_bytes, format="PNG")
                text = cls.from_image_bytes(img_bytes.getvalue(), ".png")
                if text:
                    texts.append(text)
            return "\n".join(texts)
        except Exception as e:
            log.debug(f"[OCR] PDF→image OCR failed: {e}")

        return ""

    @classmethod
    def from_docx_bytes(cls, file_bytes: bytes) -> str:
        """Extract text from DOCX — paragraphs AND table cells (critical for Aadhaar DOCX)."""
        if not DOCX_OK:
            return ""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            parts = []

            # 1. Normal paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())

            # 2. Table cells — Aadhaar DOCX stores ALL data in tables, not paragraphs
            seen = set()
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t and t not in seen:
                            seen.add(t)
                            parts.append(t)

            # 3. Headers (some DOCX put data in header)
            try:
                for section in doc.sections:
                    hdr = section.header
                    for p in hdr.paragraphs:
                        if p.text.strip():
                            parts.append(p.text.strip())
            except Exception:
                pass

            combined = "\n".join(parts)
            log.info(f"[OCR] DOCX extracted {len(combined)} chars "
                     f"({len(doc.paragraphs)} para, {len(doc.tables)} tables)")
            return combined
        except Exception as e:
            log.debug(f"[OCR] DOCX extract failed: {e}")
            return ""

    @classmethod
    def from_txt_bytes(cls, file_bytes: bytes) -> str:
        """Extract text from plain text file."""
        for enc in ["utf-8", "utf-16", "latin-1", "cp1252"]:
            try:
                return file_bytes.decode(enc)
            except Exception:
                continue
        return ""

    @classmethod
    def extract(cls, file_bytes: bytes, ext: str) -> str:
        """Main dispatcher — handles all file types."""
        ext = ext.lower()
        if ext == ".pdf":
            return cls.from_pdf_bytes(file_bytes)
        elif ext in (".docx", ".doc"):
            return cls.from_docx_bytes(file_bytes)
        elif ext in (".txt", ".text"):
            return cls.from_txt_bytes(file_bytes)
        else:
            return cls.from_image_bytes(file_bytes, ext)


# ─────────────────────────────────────────────────────────────────────────────
#  DOCUMENT PARSER — extracts structured fields from raw OCR text
# ─────────────────────────────────────────────────────────────────────────────
class DocumentParser:
    """Parses extracted OCR text into structured data for each document type."""

    # ── DOB patterns (covers all Indian ID formats) ───────────────────────
    DOB_PATTERNS = [
        # Aadhaar: DOB: 01/01/1990 or Year of Birth: 1990
        r'(?:DOB|Date\s*of\s*Birth|Birth\s*Date|D\.O\.B)[\s:]*(\d{2}[/\-\.]\d{2}[/\-\.]\d{4})',
        r'(?:DOB|Date\s*of\s*Birth)[\s:]*(\d{4})',
        # DD/MM/YYYY anywhere
        r'\b(\d{2}[/\-\.]\d{2}[/\-\.]\d{4})\b',
        # YYYY/MM/DD
        r'\b(\d{4}[/\-\.]\d{2}[/\-\.]\d{2})\b',
        # DD MMM YYYY (e.g. 01 Jan 1990)
        r'\b(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4})\b',
        # Month DD, YYYY
        r'\b((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})\b',
        # Standalone 4-digit birth year (last resort)
        r'(?:born|birth|yr|year)[\s:]*(\b(?:19[3-9]\d|200[0-9])\b)',
    ]

    # ── Name patterns ─────────────────────────────────────────────────────
    NAME_PATTERNS = [
        r'(?:Name|नाम)[\s:]+([A-Z][A-Za-z\s]{2,40})',
        r'(?:Full\s*Name|Holder\s*Name)[\s:]+([A-Z][A-Za-z\s]{2,40})',
        # Aadhaar: all caps line after "Government of India"
        r'Government\s+of\s+India\s*\n+([A-Z][A-Z\s]{4,40})',
        # Aadhaar DOCX bilingual: Tamil script followed by English name on same line
        # e.g. "சுேரஷ் ச Suresh S" or "राहुल कुमार Rahul Kumar"
        r'[\u0B80-\u0BFF\u0900-\u097F][\u0B80-\u0BFF\u0900-\u097F\s]+([A-Z][a-zA-Z\s\.]{2,35})\s*$',
        # Aadhaar: English name appears before S/O or D/O line
        r'^([A-Z][a-zA-Z\s\.]{2,35})\s*\n\s*(?:S/O|D/O|W/O|C/O)',
        # Aadhaar DOCX: English name on line before DOB line
        r'([A-Z][a-zA-Z][a-zA-Z\s\.]{1,33})\s*\n[^\n]*(?:DOB|பிறந்த|Birth)',
        # Name after "To\n" (Aadhaar letter format)
        r'(?:^|\n)To\s*\n+[^\n]*\n+([A-Z][a-zA-Z][a-zA-Z\s\.]{1,35})\n',
    ]

    # ── ID type keywords ──────────────────────────────────────────────────
    ID_KEYWORDS = {
        "Aadhaar":          ["aadhaar","aadhar","uidai","unique identification","enrollment no"],
        "PAN":              ["income tax","permanent account","pan card","govt. of india"],
        "Driving Licence":  ["driving licence","driving license","motor vehicles","dl no","licence no"],
        "Passport":         ["passport","republic of india","place of issue","nationality"],
        "Voter ID":         ["election commission","voter","epic","electors photo"],
    }

    # ── Health report keywords ────────────────────────────────────────────
    HEALTH_KEYWORDS = {
        "diagnosis":   ["diagnosis","diagnosed with","impression","findings","conclusion"],
        "doctor":      ["dr.","doctor","physician","consultant","mbbs","md,","ms,"],
        "medicines":   ["tab.","cap.","syrup","mg","injection","prescribed","rx","tablet"],
        "conditions":  {
            "Diabetes":        ["diabetes","blood sugar","fasting sugar","hba1c","insulin"],
            "Hypertension":    ["hypertension","blood pressure","bp:","mmhg"],
            "Heart Disease":   ["cardiac","coronary","heart failure","ecg","echocardiogram","angioplasty"],
            "Asthma":          ["asthma","bronchial","inhaler","spirometry","wheezing"],
            "Cancer":          ["malignancy","carcinoma","oncology","chemotherapy","biopsy"],
            "Kidney Disease":  ["kidney","renal","creatinine","dialysis","nephr"],
            "Thyroid":         ["thyroid","tsh","hyperthyroid","hypothyroid","thyroxine"],
        }
    }

    # ── Vehicle keywords ──────────────────────────────────────────────────
    VEHICLE_NO_PATTERN = r'\b([A-Z]{2}\s*\d{2}\s*[A-Z]{1,2}\s*\d{4})\b'
    VEHICLE_MODEL_KEYWORDS = [
        "maruti","honda","hyundai","toyota","ford","tata","bajaj","hero","tvs",
        "suzuki","yamaha","kia","mahindra","renault","skoda","volkswagen"
    ]

    @classmethod
    def parse_gov_id(cls, text: str) -> dict:
        """Parse a Government ID document."""
        text_lower = text.lower()
        result = {
            "id_type":       cls._detect_id_type(text_lower),
            "dob":           cls._extract_dob(text),
            "name_found":    cls._extract_name(text),
            "gender":        cls._extract_gender_from_id(text),
            "raw_text_len":  len(text),
            "text_snippet":  text[:300].replace('\n', ' '),
        }
        return result

    @classmethod
    def _extract_gender_from_id(cls, text: str) -> str:
        """
        Extract gender from Government ID text.
        Returns 'Male', 'Female', 'Other', or '' if not found.
        """
        patterns = [
            r'\b(?:Sex|Gender)\s*[:/]?\s*(Male|Female|Transgender|M|F)\b',
            r'\b(Male|Female|Transgender)\b',
            r'\bSEX\s*:\s*([MF])\b',
        ]
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                raw = m.group(1).strip().upper()
                if raw in ('M', 'MALE'):
                    return 'Male'
                elif raw in ('F', 'FEMALE'):
                    return 'Female'
                elif raw in ('TRANSGENDER', 'OTHER', 'T'):
                    return 'Other'
        return ''

    @classmethod
    def parse_health_report(cls, text: str) -> dict:
        """Parse a health/medical report."""
        text_lower = text.lower()
        found_conditions = []
        for cond, keywords in cls.HEALTH_KEYWORDS["conditions"].items():
            if any(kw in text_lower for kw in keywords):
                found_conditions.append(cond)

        # Extract doctor name
        doctor = ""
        for line in text.split('\n'):
            if any(kw in line.lower() for kw in cls.HEALTH_KEYWORDS["doctor"]):
                doctor = line.strip()[:60]
                break

        # Extract diagnosis line
        diagnosis = ""
        for line in text.split('\n'):
            if any(kw in line.lower() for kw in cls.HEALTH_KEYWORDS["diagnosis"]):
                diagnosis = line.strip()[:100]
                break

        # Extract date
        dob_result = cls._extract_dob(text)

        return {
            "conditions_found": found_conditions,
            "diagnosis":        diagnosis,
            "doctor":           doctor,
            "report_date":      dob_result.get("raw") if dob_result else None,
            "raw_text_len":     len(text),
        }

    @classmethod
    def parse_vehicle_doc(cls, text: str) -> dict:
        """Parse RC Book or previous insurance."""
        text_upper = text.upper()

        # Vehicle registration number
        vehicle_no = ""
        m = re.search(cls.VEHICLE_NO_PATTERN, text_upper)
        if m:
            vehicle_no = m.group(1).strip()

        # Model
        text_lower = text.lower()
        model_found = ""
        for brand in cls.VEHICLE_MODEL_KEYWORDS:
            if brand in text_lower:
                # Find the line with this brand
                for line in text.split('\n'):
                    if brand in line.lower():
                        model_found = line.strip()[:50]
                        break
                break

        # Year
        year_m = re.search(r'\b(20\d{2}|19[89]\d)\b', text)
        year = year_m.group(1) if year_m else ""

        # Policy number
        policy_m = re.search(r'(?:Policy\s*No|Policy\s*Number|Pol\.?\s*No)[\s:]*([A-Z0-9\-/]{6,25})', text, re.IGNORECASE)
        policy_no = policy_m.group(1).strip() if policy_m else ""

        return {
            "vehicle_number": vehicle_no,
            "vehicle_model":  model_found,
            "year":           year,
            "policy_number":  policy_no,
            "raw_text_len":   len(text),
        }

    @classmethod
    def parse_policy_doc(cls, text: str) -> dict:
        """Parse a previous insurance policy document."""
        # Policy number
        pn_m = re.search(
            r'(?:Policy\s*(?:No\.?|Number)|Pol\.\s*No)[\s:]*([A-Z0-9\-/]{6,30})',
            text, re.IGNORECASE
        )
        policy_no = pn_m.group(1).strip() if pn_m else ""

        # Coverage amount
        cov_m = re.search(
            r'(?:Sum\s*Insured|Coverage|Cover\s*Amount|Insured\s*Amount)[\s:₹]*([0-9,]+)',
            text, re.IGNORECASE
        )
        coverage = cov_m.group(1).replace(",","") if cov_m else ""

        # Premium
        prem_m = re.search(
            r'(?:Premium|Annual\s*Premium)[\s:₹]*([0-9,]+)',
            text, re.IGNORECASE
        )
        premium = prem_m.group(1).replace(",","") if prem_m else ""

        # Claim history
        has_claim = any(kw in text.lower() for kw in ["claim","settlement","claim amount","hospitalization"])

        return {
            "policy_number": policy_no,
            "coverage":      coverage,
            "premium":       premium,
            "has_claims":    has_claim,
            "raw_text_len":  len(text),
        }

    # ── Private helpers ────────────────────────────────────────────────────
    @classmethod
    def _detect_id_type(cls, text_lower: str) -> str:
        for id_type, keywords in cls.ID_KEYWORDS.items():
            if any(kw in text_lower for kw in keywords):
                return id_type
        return "Unknown"

    @classmethod
    def _extract_dob(cls, text: str) -> dict | None:
        for pattern in cls.DOB_PATTERNS:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                raw = m.group(1).strip()
                parsed = cls._parse_date(raw)
                if parsed:
                    return {"raw": raw, "parsed": parsed}
        return None

    @classmethod
    def _parse_date(cls, date_str: str) -> datetime.date | None:
        date_str = date_str.strip()
        formats = [
            "%d/%m/%Y","%d-%m-%Y","%Y-%m-%d","%d/%m/%y","%d.%m.%Y",
            "%m/%d/%Y","%d %B %Y","%d %b %Y","%B %d, %Y","%b %d, %Y",
            "%Y/%m/%d","%d %m %Y","%B %d %Y","%b %d %Y",
        ]
        for fmt in formats:
            try:
                return datetime.datetime.strptime(date_str, fmt).date()
            except ValueError:
                continue

        # Just a year?
        if re.match(r'^(19[3-9]\d|200[0-9])$', date_str):
            return datetime.date(int(date_str), 1, 1)

        return None

    @classmethod
    def _extract_name(cls, text: str) -> str:
        for pattern in cls.NAME_PATTERNS:
            m = re.search(pattern, text, re.MULTILINE)
            if m:
                name = m.group(1).strip()
                # Clean: remove extra spaces, keep only letters+spaces
                name = re.sub(r'\s+', ' ', name)
                if 2 < len(name) < 50:
                    return name
        return ""


# ─────────────────────────────────────────────────────────────────────────────
#  MAIN OCRVerifier CLASS
# ─────────────────────────────────────────────────────────────────────────────
class OCRVerifier:
    """
    Offline document verification using pytesseract + OpenCV.
    After verification completes, control passes to Gemini AI for conversation.
    """

    BAD_QUALITIES = {"blurry", "dark", "overexposed"}
    QUALITY_TIPS  = {
        "blurry":      "📸 Tip: Hold your phone steady and tap to focus before capturing.",
        "dark":        "💡 Tip: Move to a brighter area or use your phone's flashlight.",
        "overexposed": "🔆 Tip: Avoid photographing in direct sunlight or bright light sources.",
        "unknown":     "🔍 Tip: Place document flat and photograph from directly above.",
    }

    def verify_gov_id(self, file_path: str, file_bytes: bytes, file_ext: str,
                      stated_age, user_id: str, session_id: str = "",
                      stated_name: str = "") -> dict:
        """
        Full offline government ID verification pipeline.
        1. Quality check (OpenCV)
        2. OCR text extraction (pytesseract)
        3. Parse ID type + DOB + Name
        4. Compare DOB with stated age AND name with stated name
        5. Return result with handoff_to_gemini flag
        RULE: Both Name AND Age must match for verification.
        """
        sid = session_id or user_id
        print(f"[CONSOLE] Document uploaded for {sid}")
        log.info(f"[OCR-VERIFY] Starting | user={user_id} | ext={file_ext} | size={len(file_bytes)}")

        if not TESS_OK:
            return self._result("no_ocr_engine", False,
                "OCR engine not available on this server. "
                "Please ensure Tesseract is installed, or continue without verification.",
                options=["Continue Without Verification"])

        # ── Step 1: Quality check ──────────────────────────────────────────
        img = ImagePreprocessor.load_bytes(file_bytes) if file_ext not in (".pdf",".docx",".txt") else None
        quality_info = ImagePreprocessor.assess_quality(img) if img is not None else {"quality":"unknown","score":50}
        quality = quality_info["quality"]

        log.info(f"[OCR-VERIFY] Quality: {quality} | blur={quality_info.get('blur_score','?')} | brightness={quality_info.get('brightness','?')}")

        if quality in self.BAD_QUALITIES and quality_info.get("score", 100) < 30:
            tip = self.QUALITY_TIPS.get(quality, self.QUALITY_TIPS["unknown"])
            return self._result("low_quality", False,
                f"😊 Your document photo is {quality}. {tip}",
                quality=quality,
                options=["Upload Clearer Document", "Continue Without Verification"])

        # ── Step 2: OCR text extraction ────────────────────────────────────
        raw_text = TextExtractor.extract(file_bytes, file_ext)

        if not raw_text or len(raw_text.strip()) < 20:
            log.warning(f"[OCR-VERIFY] Very little text extracted: '{raw_text[:50]}'")
            return self._result("no_text", False,
                "😊 I couldn't read the text on your document. "
                "Please ensure the document is well-lit, flat, and fully in frame.",
                quality=quality,
                options=["Upload Clearer Document", "Continue Without Verification"])

        log.info(f"[OCR-VERIFY] Raw text sample: {raw_text[:200].replace(chr(10),' ')}")

        # ── Step 3: Parse ID fields ────────────────────────────────────────
        parsed = DocumentParser.parse_gov_id(raw_text)
        id_type = parsed.get("id_type", "Unknown")
        dob_info = parsed.get("dob")

        log.info(f"[OCR-VERIFY] id_type={id_type} | dob={dob_info}")

        if id_type == "Unknown":
            return self._result("not_valid_id", False,
                "😊 This doesn't appear to be a Government ID. "
                "Please upload Aadhaar, PAN, Driving Licence, Passport or Voter ID.",
                quality=quality, id_type=id_type,
                options=["Upload Government ID", "Continue Without Verification"])

        if not dob_info:
            return self._result("dob_not_found", False,
                f"😊 I could read your {id_type} but couldn't find the Date of Birth clearly. "
                "Please upload a clearer photo where all text is readable.",
                quality=quality, id_type=id_type,
                options=["Upload Clearer Document", "Continue Without Verification"])

        # ── Step 4: Age comparison ─────────────────────────────────────────
        matched, dob_year = self._compare_age(dob_info["parsed"], stated_age)

        if matched is None:
            return self._result("age_calc_error", False,
                f"😊 Found a date on your {id_type} but couldn't calculate age. "
                "Please try a different ID document.",
                quality=quality, id_type=id_type,
                options=["Upload Different ID", "Continue Without Verification"])

        if not matched:
            return self._result("age_mismatch", False,
                f"😊 The age on your {id_type} doesn't match what you told me. "
                "Please upload a different ID or continue without verification.",
                quality=quality, id_type=id_type,
                options=["Upload Different ID", "Continue Without Verification"])

        # ── Step 4b: Name comparison — STRICT: both name AND age must match ────
        name_from_id = parsed.get("name_found") or parsed.get("name") or ""
        log.info(f"[OCR-VERIFY] Name from ID: '{name_from_id}' | Stated name: '{stated_name}'")

        if stated_name and name_from_id:
            # STRICT: name on Aadhaar must match name the user gave the bot
            name_matched = self._compare_names(stated_name, name_from_id)
            log.info(f"[OCR-VERIFY] Name check: stated='{stated_name}' id='{name_from_id}' match={name_matched}")
            if not name_matched:
                # HARD REJECT — name mismatch means this is not the user's own Aadhaar
                log.warning(f"[OCR-VERIFY] HARD name mismatch: stated='{stated_name}' id='{name_from_id}'")
                return self._result("name_mismatch", False,
                    f"❌ The name on your {id_type} is '{name_from_id}', but you told me your name is '{stated_name}'. "
                    f"These don't match. Please upload YOUR OWN {id_type} that matches the name you gave me.",
                    quality=quality, id_type=id_type,
                    options=["Upload My Own ID", "Continue Without Verification"])
        elif stated_name and not name_from_id:
            # Name not found on document — let it pass (OCR may have missed it)
            log.warning(f"[OCR-VERIFY] Name not found on {id_type}, skipping name check")

        # ── Both Age AND Name verified ─────────────────────────────────────
        # Extract gender from parsed ID (never asked from user)
        id_gender = parsed.get("gender", "")
        verified_details = f"✅ Your {id_type} has been verified successfully! 👍"
        if name_from_id:
            verified_details += f" Name: {name_from_id} ✅."
        verified_details += " Identity confirmed — let's continue!"
        return self._result("verified", True,
            verified_details,
            quality=quality, id_type=id_type, dob_year=dob_year,
            name_on_id=name_from_id, id_gender=id_gender,
            handoff_to_gemini=True)

    def analyze_health_report(self, file_bytes: bytes, file_ext: str,
                               user_id: str,
                               stated_name: str = "",
                               stated_age=None,
                               stated_gender: str = "",
                               insurance_type: str = "") -> dict:
        """
        Analyze health/medical report — v2 UPGRADE.

        STEP 1 — Identity check:
          Extract patient name / age / gender from report.
          Compare with stated profile fields.
          If fields missing in report → mark as UNKNOWN (not mismatch).

        STEP 2 — Condition detection (abnormal results only):
          Do NOT flag a condition just because a test name appears.
          Only flag when:
            • status = High / Low / Abnormal / Positive
            • doctor's remark confirms the condition
            • diagnosed with / impression / conclusion line names a condition

        STEP 3 — Insurance-type filtering:
          Health → Diabetes, BP, Heart, Asthma, Kidney, Thyroid, Cancer
          Life/Term → Heart, Diabetes, Cancer, Major illness
          Accident → skip medical (return healthy)
        """
        log.info(f"[OCR-HEALTH] Analyzing health report | user={user_id} "
                 f"| name='{stated_name}' age={stated_age} gender='{stated_gender}'")

        raw_text = TextExtractor.extract(file_bytes, file_ext)
        if not raw_text or len(raw_text.strip()) < 10:
            return {
                "success":  False,
                "message":  "Could not read the health report. Please upload a clearer document.",
                "conditions": [],
                "identity_check": {"name": "UNKNOWN", "age": "UNKNOWN", "gender": "UNKNOWN"},
            }

        text_lower = raw_text.lower()

        # ══════════════════════════════════════════════════════════════════════
        # STEP 1 — Identity check from report header
        # ══════════════════════════════════════════════════════════════════════
        identity = self._extract_report_identity(raw_text)
        id_check = {
            "name":   "UNKNOWN",
            "age":    "UNKNOWN",
            "gender": "UNKNOWN",
        }
        id_warnings = []

        # Name check
        if stated_name and identity.get("name"):
            matched = self._compare_names(stated_name, identity["name"])
            id_check["name"] = "YES" if matched else "NO"
            if not matched:
                id_warnings.append(
                    f"⚠️ Patient name on report ('{identity['name']}') "
                    f"doesn't match your name ('{stated_name}')."
                )
        elif identity.get("name"):
            id_check["name"] = "UNKNOWN"   # no stated name to compare

        # Age check
        if stated_age and identity.get("age") is not None:
            try:
                doc_age = int(identity["age"])
                diff    = abs(doc_age - int(stated_age))
                id_check["age"] = "YES" if diff <= 3 else "NO"
                if diff > 3:
                    id_warnings.append(
                        f"⚠️ Patient age on report ({doc_age}) "
                        f"doesn't match your age ({stated_age})."
                    )
            except (ValueError, TypeError):
                id_check["age"] = "UNKNOWN"

        # Gender check
        if stated_gender and identity.get("gender"):
            doc_gender = identity["gender"].lower()
            st_gender  = stated_gender.lower()
            # Normalize to M/F
            is_male   = st_gender in ("male", "m")
            is_female = st_gender in ("female", "f")
            doc_male   = any(w in doc_gender for w in ("male", " m ", "m/"))
            doc_female = any(w in doc_gender for w in ("female", " f ", "f/"))
            if (is_male and doc_male) or (is_female and doc_female):
                id_check["gender"] = "YES"
            elif doc_male or doc_female:
                id_check["gender"] = "NO"
                id_warnings.append(
                    f"⚠️ Gender on report ('{identity['gender']}') "
                    f"doesn't match your profile ('{stated_gender}')."
                )

        log.info(f"[OCR-HEALTH] Identity check: {id_check} | warnings={id_warnings}")

        # ══════════════════════════════════════════════════════════════════════
        # STEP 2 — Condition detection: abnormal results ONLY
        # ══════════════════════════════════════════════════════════════════════
        ins_lower = (insurance_type or "").lower()

        # Accident insurance → no medical check needed
        if "accident" in ins_lower:
            parsed     = DocumentParser.parse_health_report(raw_text)
            doctor     = parsed.get("doctor", "")
            conditions = []
            risk_level = "LOW"
        else:
            conditions = self._detect_abnormal_conditions(raw_text, text_lower)
            parsed     = DocumentParser.parse_health_report(raw_text)
            doctor     = parsed.get("doctor", "")
            diagnosis  = parsed.get("diagnosis", "")

            # Insurance-type filter: only keep relevant conditions
            if "life" in ins_lower or "term" in ins_lower:
                relevant = {"Heart Disease", "Diabetes", "Cancer", "Kidney Disease",
                            "Hypertension", "Thyroid"}
                conditions = [c for c in conditions if c in relevant]
            # Health → all conditions relevant (no filter needed)
            # Travel → similar to health

            risk_level = ("HIGH"   if len(conditions) >= 2 else
                          "MEDIUM" if len(conditions) == 1 else
                          "LOW")

        # ══════════════════════════════════════════════════════════════════════
        # STEP 3 — Build result message
        # ══════════════════════════════════════════════════════════════════════
        lines_out = ["📋 **Document Verification Result**", ""]
        lines_out.append(f"**Document Type:** Medical / Health Report")
        lines_out.append("")
        lines_out.append("**Identity Check**")
        lines_out.append(f"• Name Match: {id_check['name']}")
        lines_out.append(f"• Age Match:  {id_check['age']}")
        lines_out.append(f"• Gender Match: {id_check['gender']}")

        if id_warnings:
            lines_out.append("")
            lines_out.extend(id_warnings)

        lines_out.append("")
        lines_out.append("**Health Analysis**")

        if conditions:
            lines_out.append(f"• Detected Conditions: {', '.join(conditions)}")
        else:
            lines_out.append("• Conclusion: All parameters are normal. No medical conditions detected.")

        if doctor:
            lines_out.append(f"• Issuing Doctor / Lab: {doctor}")

        lines_out.append(f"\n**Insurance Risk Level: {risk_level}**")
        lines_out.append("")
        lines_out.append("I'll factor this into your plan recommendation 👍")

        reply_msg = "\n".join(lines_out)
        log.info(f"[OCR-HEALTH] Conditions={conditions} | Risk={risk_level}")

        return {
            "success":       True,
            "message":       reply_msg,
            "conditions":    conditions,
            "doctor":        doctor,
            "identity_check": id_check,
            "id_warnings":   id_warnings,
            "risk_level":    risk_level,
            "handoff_to_gemini": True,
        }

    # ── Identity extraction from medical report header ─────────────────────
    def _extract_report_identity(self, text: str) -> dict:
        """Extract patient name, age, gender from medical report header."""
        result = {}

        # Patient Name — common labels in lab reports
        name_patterns = [
            r'(?:Patient\s*Name|Name)[\\s:]+([A-Za-z][A-Za-z\\s\\.]{2,40})',
            r'(?:Mr\\.?|Mrs\\.?|Ms\\.?|Dr\\.?)\\s+([A-Z][A-Za-z\\s]{2,35})',
            r'(?:Ref\\.?\\s*by|Referred\\s*by)[\\s:]+.{0,20}\\n([A-Z][A-Za-z\\s]{3,35})',
        ]
        for pat in name_patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                candidate = m.group(1).strip()
                candidate = re.sub(r'\s+', ' ', candidate)
                if 3 < len(candidate) < 50:
                    result["name"] = candidate
                    break

        # Age — common formats: Age: 34Y, Age/Sex: 34Y/M, Age: 34 Years
        age_patterns = [
            r'Age[\s:/]*(\d{1,3})\s*(?:Y(?:rs?|ears?)?|year)',
            r'Age\s*/\s*Sex[\s:]*(\d{1,3})',
            r'(\d{1,3})\s*(?:Years?|Yrs?|Y)\s*/\s*(?:M|F|Male|Female)',
        ]
        for pat in age_patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                age_val = int(m.group(1))
                if 1 <= age_val <= 120:
                    result["age"] = age_val
                    break

        # Gender
        gender_pat = r'(?:Sex|Gender)[\s:/]*([MF](?:ale)?|Male|Female)'
        m = re.search(gender_pat, text, re.IGNORECASE)
        if m:
            result["gender"] = m.group(1).strip()

        return result

    # ── Abnormal-only condition detector ──────────────────────────────────
    def _detect_abnormal_conditions(self, raw_text: str, text_lower: str) -> list:
        """
        SMART DETECTION: only flag a condition when results are abnormal.

        Logic:
          • Look for the condition keyword AND an abnormal marker nearby
            (HIGH / LOW / ABNORMAL / POSITIVE / ELEVATED / ABOVE NORMAL)
          • OR look for explicit diagnosis/impression lines naming the disease
          • Do NOT flag just because a test name appears with a normal result
        """
        ABNORMAL_MARKERS = [
            "high", "low", "abnormal", "positive", "elevated",
            "above normal", "below normal", "increased", "decreased",
            "borderline", "critical", "flagged", "out of range",
            "uncontrolled", "detected",
        ]

        # Tests that are NOT diseases by themselves
        TEST_ONLY_TERMS = {
            "blood sugar", "fbs", "rbs", "ppbs", "hba1c",    # glucose tests (not disease)
            "cholesterol", "ldl", "hdl", "triglycerides",
            "tsh", "t3", "t4",                                # thyroid tests
            "creatinine", "egfr", "urea",                     # kidney tests
            "bilirubin", "sgot", "sgpt", "alt", "ast",        # liver tests
            "haemoglobin", "hb", "rbc", "wbc", "platelets",
        }

        DIAGNOSIS_MARKERS = [
            "diagnosis", "diagnosed with", "impression", "conclusion",
            "findings", "assessment", "final report", "clinical diagnosis",
        ]

        def window_has_abnormal(text: str, pos: int, window: int = 120) -> bool:
            """Check if any abnormal marker appears within ±window chars of pos."""
            start = max(0, pos - window)
            end   = min(len(text), pos + window)
            snippet = text[start:end].lower()
            return any(m in snippet for m in ABNORMAL_MARKERS)

        # ── Condition map: keyword → condition label ───────────────────────
        CONDITION_MAP = {
            "Diabetes": [
                "diabetes", "diabetic", "insulin", "blood sugar high",
                "hba1c high", "fasting sugar high", "hyperglycemia",
            ],
            "Hypertension": [
                "hypertension", "blood pressure high", "bp elevated",
                "systolic high", "diastolic high",
            ],
            "Heart Disease": [
                "coronary", "cardiac", "heart failure", "myocardial",
                "angina", "arrhythmia", "ecg abnormal", "angioplasty",
                "left ventricular", "right ventricular",
            ],
            "Asthma": [
                "asthma", "bronchial asthma", "inhaler", "wheezing abnormal",
                "spirometry abnormal", "obstructive airway",
            ],
            "Cancer": [
                "malignancy", "carcinoma", "tumor", "oncology",
                "chemotherapy", "biopsy positive", "metastasis",
            ],
            "Kidney Disease": [
                "chronic kidney", "renal failure", "creatinine high",
                "dialysis", "nephropathy", "kidney disease",
                "gfr low", "proteinuria",
            ],
            "Thyroid": [
                "hypothyroid", "hyperthyroid", "tsh high", "tsh low",
                "tsh elevated", "thyroid disorder", "thyroiditis",
            ],
            "Liver Disease": [
                "hepatitis", "cirrhosis", "liver failure", "sgot high",
                "sgpt high", "alt elevated", "fatty liver grade",
            ],
        }

        found_conditions = []

        # ── Pass 1: keyword + abnormal marker proximity ────────────────────
        for condition, keywords in CONDITION_MAP.items():
            if condition in found_conditions:
                continue
            for kw in keywords:
                pos = text_lower.find(kw)
                if pos == -1:
                    continue
                # Skip if this is a test-only term and no abnormal nearby
                if kw in TEST_ONLY_TERMS:
                    if window_has_abnormal(text_lower, pos):
                        found_conditions.append(condition)
                        break
                else:
                    # Non-test keyword found → immediate flag (e.g. "coronary", "cirrhosis")
                    found_conditions.append(condition)
                    break

        # ── Pass 2: scan diagnosis / impression lines ──────────────────────
        for line in raw_text.split('\n'):
            line_l = line.lower().strip()
            if not any(dm in line_l for dm in DIAGNOSIS_MARKERS):
                continue
            # This is a diagnosis line — check all conditions
            for condition, keywords in CONDITION_MAP.items():
                if condition in found_conditions:
                    continue
                if any(kw in line_l for kw in keywords):
                    found_conditions.append(condition)

        return list(dict.fromkeys(found_conditions))   # preserve order, deduplicate

    def analyze_vehicle_doc(self, file_bytes: bytes, file_ext: str,
                             user_id: str) -> dict:
        """Analyze RC Book or vehicle insurance."""
        raw_text = TextExtractor.extract(file_bytes, file_ext)
        if not raw_text or len(raw_text.strip()) < 10:
            return {"success": False, "message": "Could not read the vehicle document.", "vehicle_no": ""}

        parsed = DocumentParser.parse_vehicle_doc(raw_text)
        vehicle_no = parsed.get("vehicle_number", "")
        model      = parsed.get("vehicle_model", "")
        year       = parsed.get("year", "")

        msg = "✅ Vehicle document analyzed! "
        if vehicle_no: msg += f"Vehicle: {vehicle_no}. "
        if model:      msg += f"Model: {model}. "
        if year:       msg += f"Year: {year}. "
        msg += "Great, let's continue 👍"

        return {
            "success":      True,
            "message":      msg,
            "vehicle_no":   vehicle_no,
            "vehicle_model": model,
            "year":         year,
            "handoff_to_gemini": True,
        }

    def analyze_policy_doc(self, file_bytes: bytes, file_ext: str,
                            user_id: str) -> dict:
        """Analyze previous insurance policy document."""
        raw_text = TextExtractor.extract(file_bytes, file_ext)
        if not raw_text or len(raw_text.strip()) < 10:
            return {"success": False, "message": "Could not read the policy document.", "policy_no": ""}

        parsed = DocumentParser.parse_policy_doc(raw_text)

        msg = "✅ Previous policy analyzed! "
        if parsed.get("policy_number"): msg += f"Policy No: {parsed['policy_number']}. "
        if parsed.get("coverage"):      msg += f"Coverage: ₹{parsed['coverage']}. "
        if parsed.get("has_claims"):    msg += "Claim history found — noted for recommendation. "
        msg += "This helps me suggest better plans 👍"

        return {
            "success":      True,
            "message":      msg,
            "policy_no":    parsed.get("policy_number",""),
            "coverage":     parsed.get("coverage",""),
            "has_claims":   parsed.get("has_claims", False),
            "handoff_to_gemini": True,
        }

    def extract_policy_text_for_rag(self, file_bytes: bytes, file_ext: str) -> str:
        """Extract full text from a policy PDF for RAG embedding."""
        return TextExtractor.extract(file_bytes, file_ext)

    # ── Helpers ────────────────────────────────────────────────────────────
    def _compare_names(self, stated: str, from_id: str) -> bool:
        """
        Strict fuzzy name match.
        Rules:
          1. Exact word match (case-insensitive) — "Suresh" == "Suresh" ✅
          2. Substring prefix — stated word must be a REAL prefix of the ID word AND
             cover >= 85% of the ID word length.
             e.g. "sures" (5) vs "suresh" (6): 5/6 = 83% → FAIL (below 85%)
             e.g. "suresh" (6) vs "suresh" (6): exact → PASS
             e.g. "surya" (5) vs "suryakumar" (10): 5/10 = 50% → FAIL
          3. If either side is empty after cleaning → can't compare → benefit of doubt.
        Fail examples:  "sures"  vs "Suresh S"   → ❌ (83% < 85%)
                        "Mithun" vs "Suresh S"   → ❌ (zero overlap)
        Pass examples:  "Suresh" vs "Suresh S"   → ✅ (exact)
                        "suresh" vs "Suresh"     → ✅ (exact, lowercased)
        """
        import re as _re
        stop = {"mr","mrs","ms","dr","shri","smt","kumari","late","s/o","d/o","w/o","c/o"}
        def clean(s):
            s = _re.sub(r"[^a-z ]", "", s.strip().lower())
            return [w for w in s.split() if len(w) >= 2 and w not in stop]
        s_words = clean(stated)
        d_words = clean(from_id)
        if not d_words:
            return True   # can't extract name from ID — OCR missed it — benefit of doubt
        if not s_words:
            return False  # stated name is too short/incomplete to verify
        s_set = set(s_words)
        d_set = set(d_words)
        # Rule 1: exact word match
        if s_set & d_set:
            return True
        # Rule 2: real substring prefix — stated word must be actual prefix of ID word
        # AND must cover >= 85% of ID word (prevents "sures" matching "suresh" at 83%)
        for sw in s_words:
            if len(sw) < 4:   # too short to prefix-match reliably
                continue
            for dw in d_words:
                if len(dw) < 4:
                    continue
                # sw must be a true prefix of dw (or vice-versa) with >=85% coverage
                if dw.startswith(sw) and len(sw) / len(dw) >= 0.88:
                    return True
                if sw.startswith(dw) and len(dw) / len(sw) >= 0.88:
                    return True
        return False

    def _compare_age(self, dob: datetime.date | None, stated_age) -> tuple:
        if dob is None or stated_age is None:
            return None, ""
        try:
            calc_age = (datetime.date.today() - dob).days // 365
            match = abs(calc_age - int(stated_age)) <= 3   # ±3 yr tolerance (accounts for birthday not yet passed)
            return match, str(dob.year)
        except (ValueError, TypeError):
            return None, ""

    def _result(self, status, verified, message, quality="", id_type="",
                dob_year="", notes="", options=None, handoff_to_gemini=False,
                name_on_id="", id_gender="") -> dict:
        return {
            "status":           status,
            "verified":         verified,
            "message":          message,
            "quality":          quality,
            "doc_type_found":   id_type,
            "dob_year":         dob_year,
            "notes":            notes,
            "options":          options or [],
            "option_type":      "radio" if options else "none",
            "handoff_to_gemini": handoff_to_gemini,
            "engine":           "offline_ocr",
            "name_on_id":       name_on_id,
            "id_gender":        id_gender,   # extracted from ID, never asked manually
        }


# ─────────────────────────────────────────────────────────────────────────────
#  AVAILABILITY CHECK
# ─────────────────────────────────────────────────────────────────────────────
def ocr_available() -> dict:
    """Check what OCR components are available on this system."""
    return {
        "tesseract":  TESS_OK,
        "opencv":     CV2_OK,
        "pil":        PIL_OK,
        "pdf_read":   PDF_OK,
        "docx_read":  DOCX_OK,
        "fully_ready": TESS_OK and PIL_OK,
    }
