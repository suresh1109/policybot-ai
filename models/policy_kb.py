"""
PolicyBot — Policy Knowledge Base Engine
Handles: upload → OCR/parse → AI extraction → DB storage → recommendation
"""
import os, re, json, logging, hashlib
log = logging.getLogger("PolicyBot.KB")

# ── Optional imports (graceful degradation) ───────────────────────────────────
try:
    import pdfplumber; PDFPLUMBER_OK = True
except ImportError:
    PDFPLUMBER_OK = False

try:
    import pypdf; PYPDF_OK = True
except ImportError:
    PYPDF_OK = False

# Legacy PyPDF2 fallback
try:
    import PyPDF2; PYPDF2_OK = True
except ImportError:
    PYPDF2_OK = False

PDF_OK = PDFPLUMBER_OK or PYPDF_OK or PYPDF2_OK

try:
    import docx as docx_lib; DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import pytesseract
    from PIL import Image
    import io as _io
    OCR_OK = True
except ImportError:
    OCR_OK = False

ALLOWED_EXTS = {".pdf", ".txt", ".docx", ".doc", ".jpg", ".jpeg", ".png", ".webp"}

EXTRACTION_PROMPT = """You are an expert Indian insurance policy data extractor.
Read the document text below and extract EVERY insurance plan found.

This may be a CATALOG with MANY plans across multiple sections and tables.
You MUST extract EVERY SINGLE PLAN — do not stop early.

Return ONLY a valid JSON array. Each element is one insurance plan with EXACTLY these keys:
  company_name, plan_name, insurance_type, coverage_amount, premium_range,
  waiting_period, conditions_covered, exclusions, claim_process,
  network_hospitals, eligibility_age, special_benefits, raw_summary

EXTRACTION RULES:
- insurance_type: MUST be one of: Health Insurance, Life Insurance, Term Life Insurance,
  Vehicle Insurance, Travel Insurance, Property Insurance, Accident Insurance
- company_name: The insurer name (Star Health, HDFC ERGO, Bajaj Allianz, ICICI Lombard,
  LIC India, SBI Life, Tata AIG, Niva Bupa, Care Health, Aditya Birla, New India,
  United India, Oriental Insurance, Reliance General, Digit Insurance, ICICI Prudential,
  Max Life, Kotak, Royal Sundaram, Manipal Cigna, National Insurance, etc.)
- plan_name: The exact product name as written (e.g. "Optima Restore", "Click2Protect Super",
  "Senior Citizen Red Carpet", "Personal Guard", "Travel Assurance International")
- coverage_amount: Sum insured / coverage limit (e.g. "Rs.5 Lakh to Rs.1 Crore")
- premium_range: Annual premium (e.g. "Rs.8,000-Rs.25,000/year")
- waiting_period: PED or initial waiting (e.g. "30 days general, 2 yrs PED"; "Not specified" if absent)
- conditions_covered: Diseases/events covered (e.g. "Hospitalisation, Surgery, ICU, Daycare")
- exclusions: What is NOT covered (e.g. "Cosmetic surgery, self-inflicted injuries")
- claim_process: How to claim (e.g. "Cashless at network hospitals or reimbursement")
- network_hospitals: Network size (e.g. "10,000+ network hospitals"; "Not specified" if absent)
- eligibility_age: Age range (e.g. "18-65 years"; "Not specified" if absent)
- special_benefits: Unique features (e.g. "No-claim bonus 50%, free health check")
- raw_summary: 1-2 sentence plain English summary of this specific plan

HOW TO READ TABLE ROWS IN THIS DOCUMENT:
- Lines starting with "HEADER:" list the column names for the following rows
- All other pipe-separated lines ( value1 | value2 | value3 ) are data rows
- EACH data row in a plan table = ONE separate plan in your JSON output
- Map each pipe-separated value to the column name above it

EXAMPLE — if you see:
  HEADER: Plan Name | Insurer | Sum Insured | Annual Premium | Key Feature | Best For
  Arogya Sanjeevani | All Insurers | Rs.1L-Rs.5L | Rs.3,000-Rs.8,000 | Standard policy | Budget buyers
  MyHealth Suraksha | HDFC ERGO | Rs.3L-Rs.75L | Rs.5,500-Rs.22,000 | No room rent cap | Mid-income

You MUST output TWO plan objects — one for Arogya Sanjeevani and one for MyHealth Suraksha.

IMPORTANT:
- A document with 8 tables of 8-10 rows each = ~80 plans — extract ALL of them
- If a field is not in the document for a plan, write exactly: "Not specified"
- Do NOT merge multiple plans into one entry
- Return ONLY the JSON array — no preamble, no explanation, no markdown fences

Document text:
{text}
"""

SECURITY_PATTERNS = [
    b"<script", b"javascript:", b"<?php", b"eval(",
    b"exec(", b"system(", b"__import__",
]

class PolicyKB:
    def __init__(self, db, gemini, ocr_verifier=None):
        self.db     = db
        self.gemini = gemini
        self.ocr    = ocr_verifier

    # ── File security scan ────────────────────────────────────────────────────
    def _scan_file(self, file_bytes: bytes, ext: str) -> tuple[bool, str]:
        """Basic security scan — check for embedded scripts."""
        lower = file_bytes[:4096].lower()
        for pat in SECURITY_PATTERNS:
            if pat in lower:
                return False, f"File rejected: suspicious pattern '{pat.decode()}' detected"
        if len(file_bytes) > 25 * 1024 * 1024:
            return False, "File too large. Maximum 25MB for policy documents."
        if len(file_bytes) < 100:
            return False, "File appears empty or corrupted."
        return True, "OK"

    # ── Text extraction ───────────────────────────────────────────────────────
    def _extract_text(self, file_bytes: bytes, ext: str, file_path: str) -> str:
        text = ""
        try:
            if ext == ".pdf" and PDF_OK:
                import io
                text = ""

                # Strategy 1: pdfplumber (best for tables, columns, complex layouts)
                if PDFPLUMBER_OK:
                    try:
                        import pdfplumber
                        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                            pages_text = []
                            for page in pdf.pages:
                                # Extract regular text
                                page_text = page.extract_text() or ""
                                # Also extract tables as text
                                tables = page.extract_tables()
                                for table in tables:
                                    for row in table:
                                        row_clean = [str(c or "").strip() for c in row]
                                        pages_text.append(" | ".join(row_clean))
                                if page_text:
                                    pages_text.append(page_text)
                            text = "\n".join(pages_text)
                        log.info(f"[KB] pdfplumber: {len(text)} chars from {len(pdf.pages)} pages")
                    except Exception as e:
                        log.warning(f"[KB] pdfplumber failed: {e}")
                        text = ""

                # Strategy 2: pypdf (fast, good for text-based PDFs)
                if not text.strip() and PYPDF_OK:
                    try:
                        import pypdf
                        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                        text = "\n".join(
                            page.extract_text() or "" for page in reader.pages
                        )
                        log.info(f"[KB] pypdf: {len(text)} chars from {len(reader.pages)} pages")
                    except Exception as e:
                        log.warning(f"[KB] pypdf failed: {e}")
                        text = ""

                # Strategy 3: PyPDF2 legacy fallback
                if not text.strip() and PYPDF2_OK:
                    try:
                        import PyPDF2
                        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                        text = "\n".join(
                            page.extract_text() or "" for page in reader.pages
                        )
                        log.info(f"[KB] PyPDF2: {len(text)} chars")
                    except Exception as e:
                        log.warning(f"[KB] PyPDF2 failed: {e}")
                        text = ""

                if text.strip():
                    log.info(f"[KB] PDF total: {len(text)} chars extracted")

            elif ext in (".docx", ".doc") and DOCX_OK:
                import docx as docx_lib, io
                from docx.oxml.ns import qn as _qn
                doc = docx_lib.Document(io.BytesIO(file_bytes))
                parts = []
                # Walk body in reading order — paragraphs AND tables interleaved.
                # doc.paragraphs alone silently skips ALL table content, so a
                # document whose plans are in tables returns zero plan data.
                for child in doc.element.body:
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag == 'p':
                        para_text = ''.join(
                            r.text for r in child.iter(_qn('w:t'))
                        ).strip()
                        if para_text:
                            parts.append(para_text)
                    elif tag == 'tbl':
                        rows = child.findall('.//' + _qn('w:tr'))
                        for ri, row in enumerate(rows):
                            cells_text = []
                            for cell in row.findall('.//' + _qn('w:tc')):
                                cell_val = ' '.join(
                                    ''.join(r2.text for r2 in p2.iter(_qn('w:t'))).strip()
                                    for p2 in cell.findall('.//' + _qn('w:p'))
                                ).strip()
                                if cell_val:
                                    cells_text.append(cell_val)
                            if cells_text:
                                prefix = 'HEADER: ' if ri == 0 else ''
                                parts.append(prefix + ' | '.join(cells_text))
                text = '\n'.join(parts)
                log.info(f"[KB] DOCX: {len(text)} chars ({len(doc.paragraphs)} paras + {len(doc.tables)} tables)")

            elif ext == ".txt":
                for enc in ("utf-8", "utf-16", "latin-1"):
                    try:
                        text = file_bytes.decode(enc)
                        break
                    except Exception:
                        continue
                log.info(f"[KB] TXT: {len(text)} chars")

            elif ext in (".jpg", ".jpeg", ".png", ".webp") and OCR_OK:
                from PIL import Image
                import io
                img = Image.open(io.BytesIO(file_bytes))
                text = pytesseract.image_to_string(img)
                log.info(f"[KB] OCR image: {len(text)} chars")

        except Exception as e:
            log.error(f"[KB] Text extraction error ({ext}): {e}")

        # Fallback 1: Tesseract OCR via pdf2image
        if not text.strip() and OCR_OK and ext == ".pdf":
            try:
                from pdf2image import convert_from_bytes
                images = convert_from_bytes(file_bytes, dpi=200)
                for img in images[:10]:
                    text += pytesseract.image_to_string(img) + "\n"
                log.info(f"[KB] PDF-OCR fallback: {len(text)} chars")
            except Exception as e:
                log.warning(f"[KB] PDF-OCR fallback failed: {e}")

        # Fallback 2: Gemini Vision — converts PDF pages to images inline
        # Works for scanned PDFs, image-heavy brochures, complex layouts
        if not text.strip() and ext == ".pdf" and self.gemini:
            try:
                import io, base64
                text = self._extract_text_via_gemini_vision(file_bytes)
                if text:
                    log.info(f"[KB] Gemini Vision PDF fallback: {len(text)} chars")
            except Exception as e:
                log.warning(f"[KB] Gemini Vision fallback failed: {e}")

        return text.strip()

    def _extract_text_via_gemini_vision(self, file_bytes: bytes) -> str:
        """Use Gemini Vision to extract text from scanned/image PDFs.
        Converts pages to PNG and sends to Gemini for OCR + structure extraction."""
        try:
            import io, base64

            # Try to render PDF pages as images using pdfplumber's internal rendering
            # Or use pypdf to detect if it's a scanned (image-only) PDF
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            num_pages = len(reader.pages)

            # Check if PDF has actual text — if it does, Vision is unnecessary
            sample_text = "".join(
                (reader.pages[i].extract_text() or "") for i in range(min(3, num_pages))
            )
            if len(sample_text.strip()) > 100:
                return sample_text  # Has real text, use it

            # Scanned PDF — try to use pdf2image if available
            try:
                from pdf2image import convert_from_bytes
                from PIL import Image
                pages = convert_from_bytes(file_bytes, dpi=150, fmt="PNG")
                all_text = []
                for i, page_img in enumerate(pages[:8]):  # Max 8 pages
                    buf = io.BytesIO()
                    page_img.save(buf, format="PNG")
                    img_b64 = base64.b64encode(buf.getvalue()).decode()
                    vision_prompt = (
                        f"This is page {i+1} of an Indian insurance policy document. "
                        "Extract ALL text exactly as it appears. Include tables, numbers, "
                        "plan names, coverage amounts, premiums, conditions, exclusions. "
                        "Return only the extracted text, no commentary."
                    )
                    # Use Gemini's multimodal capability
                    page_text = self.gemini.generate_with_image(vision_prompt, img_b64, "image/png")
                    if page_text:
                        all_text.append(page_text)
                return "\n\n".join(all_text)
            except ImportError:
                # pdf2image not available — try direct base64 PDF to Gemini
                pdf_b64 = base64.b64encode(file_bytes).decode()
                vision_prompt = (
                    "This is an Indian insurance policy document PDF. "
                    "Extract ALL plan information: plan names, company names, coverage amounts, "
                    "premiums, waiting periods, conditions covered, exclusions, benefits. "
                    "Return the extracted information as structured text."
                )
                return self.gemini.generate_with_pdf(vision_prompt, pdf_b64) or ""
        except Exception as e:
            log.warning(f"[KB] _extract_text_via_gemini_vision error: {e}")
            return ""

    # ── AI structured extraction ──────────────────────────────────────────────
    def _extract_plans_via_ai(self, text: str) -> list:
        """Use Gemini to extract structured plan data from raw text.
        Handles large documents by processing in chunks.
        Has a direct parser for structured PLAN XX / key: value format."""
        if not text:
            return []

        # ── DIRECT PARSER — runs BEFORE Gemini, no API call needed ───────────
        # Works on documents using the structured PLAN 01 / key: value format.
        # 100% reliable, instant, no quota usage.
        FIELDS = ["plan_name","company_name","insurance_type","coverage_amount",
                  "premium_range","waiting_period","eligibility_age",
                  "conditions_covered","exclusions","claim_process",
                  "network_hospitals","special_benefits","raw_summary"]
        plan_blocks = re.split(r'\nPLAN \d+\n-{10,}\n', text)
        if len(plan_blocks) > 1:
            direct_plans = []
            for block in plan_blocks[1:]:
                plan = {}
                for field in FIELDS:
                    m = re.search(rf'^{re.escape(field)}:\s*(.+)$', block, re.MULTILINE)
                    plan[field] = m.group(1).strip() if m else "Not specified"
                if plan.get("plan_name", "Not specified") != "Not specified":
                    direct_plans.append(plan)
            if direct_plans:
                log.info(f"[KB] Direct parser: {len(direct_plans)} plan(s) extracted (no AI needed)")
                return direct_plans

        # ── GEMINI AI EXTRACTION (fallback for unstructured documents) ────────
        # Clean text — remove excessive whitespace and null chars
        text = re.sub(r'\x00', '', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # For large documents, process in overlapping chunks to catch all plans
        # 18000-char chunks with 7000 output tokens handles 20-25 plans per chunk
        MAX_CHUNK  = 18000
        MAX_TOKENS = 7000
        all_plans = []
        seen_names = set()

        chunks = []
        if len(text) <= MAX_CHUNK:
            chunks = [text]
        else:
            # Overlap chunks by 500 chars so plans at boundaries aren't missed
            step = MAX_CHUNK - 500
            for i in range(0, len(text), step):
                chunks.append(text[i:i+MAX_CHUNK])
                if i + MAX_CHUNK >= len(text):
                    break

        log.info(f"[KB] Processing {len(chunks)} chunk(s) from {len(text)} chars")

        for chunk_idx, chunk in enumerate(chunks):
            prompt = EXTRACTION_PROMPT.format(text=chunk)
            try:
                response = self.gemini.generate(prompt, max_tokens=MAX_TOKENS)
                if not response:
                    continue

                # Strip markdown code fences
                raw = response.strip()
                raw = re.sub(r'^```json\s*', '', raw, flags=re.MULTILINE)
                raw = re.sub(r'\s*```\s*$', '', raw, flags=re.MULTILINE)
                raw = raw.strip()

                # Find JSON array
                start = raw.find('[')
                end   = raw.rfind(']') + 1
                if start == -1 or end == 0:
                    log.warning(f"[KB] Chunk {chunk_idx}: no JSON array found")
                    continue

                chunk_plans = json.loads(raw[start:end])
                if isinstance(chunk_plans, list):
                    for plan in chunk_plans:
                        pname = (plan.get("plan_name") or "").strip()
                        if pname and pname not in seen_names and pname != "Not specified":
                            seen_names.add(pname)
                            all_plans.append(plan)
                    log.info(f"[KB] Chunk {chunk_idx}: {len(chunk_plans)} plan(s) extracted")

            except json.JSONDecodeError as e:
                log.warning(f"[KB] Chunk {chunk_idx} JSON error: {e}")
                continue
            except Exception as e:
                log.error(f"[KB] Chunk {chunk_idx} error: {e}")
                continue

        if all_plans:
            log.info(f"[KB] Total: {len(all_plans)} unique plan(s) extracted")
            return all_plans

        log.warning("[KB] AI extraction yielded no plans, using fallback")
        return self._fallback_extract(text)

    def _fallback_extract(self, text: str) -> list:
        """Regex-based extraction when AI unavailable or fails."""
        plans = []
        text_lower = text.lower()

        # Detect insurance type
        ins_type = "Health Insurance"
        if any(w in text_lower for w in ["term life", "term insurance", "jeevan"]):
            ins_type = "Term Life Insurance"
        elif any(w in text_lower for w in ["life insurance", "whole life", "endowment", "ulip"]):
            ins_type = "Life Insurance"
        elif any(w in text_lower for w in ["vehicle", "motor", "car insurance", "bike insurance", "two-wheeler"]):
            ins_type = "Vehicle Insurance"
        elif any(w in text_lower for w in ["travel insurance", "overseas", "trip"]):
            ins_type = "Travel Insurance"
        elif any(w in text_lower for w in ["property", "home insurance", "house", "building"]):
            ins_type = "Property Insurance"
        elif any(w in text_lower for w in ["personal accident", "accidental death"]):
            ins_type = "Accident Insurance"

        # Extract coverage amount
        coverage = "Not specified"
        m = re.search(r'(?:sum insured|coverage|cover)\s*[:\-]?\s*(₹[\d,]+\s*(?:lakh|crore|L|Cr)?(?:\s*[-–to]+\s*₹[\d,]+\s*(?:lakh|crore|L|Cr)?)?)', text, re.I)
        if m:
            coverage = m.group(1).strip()

        # Extract premium
        premium = "Not specified"
        m = re.search(r'premium\s*[:\-]?\s*(₹[\d,]+\s*(?:per month|\/month|p\.m\.)?(?:\s*[-–to]+\s*₹[\d,]+)?)', text, re.I)
        if m:
            premium = m.group(1).strip()

        # Extract waiting period
        waiting = "Not specified"
        m = re.search(r'waiting period\s*[:\-]?\s*([^\n.]{5,60})', text, re.I)
        if m:
            waiting = m.group(1).strip()

        # Guess company and plan name from first 300 chars
        header = text[:300]
        company = "Not specified"
        for c in ["Star Health", "HDFC ERGO", "Niva Bupa", "Care Health", "Bajaj Allianz",
                  "ICICI Lombard", "New India", "LIC", "SBI Life", "Tata AIG", "Aditya Birla"]:
            if c.lower() in header.lower():
                company = c
                break

        plan_name = "Unknown Plan"
        m = re.search(r'plan\s+name\s*[:\-]?\s*([^\n]{3,60})', header, re.I)
        if m:
            plan_name = m.group(1).strip()
        elif company != "Not specified":
            plan_name = f"{company} Policy"

        plans.append({
            "company_name":       company,
            "plan_name":          plan_name,
            "insurance_type":     ins_type,
            "coverage_amount":    coverage,
            "premium_range":      premium,
            "waiting_period":     waiting,
            "conditions_covered": "As per policy document",
            "exclusions":         "As per policy document",
            "claim_process":      "Contact insurer helpline",
            "network_hospitals":  "As per insurer network",
            "eligibility_age":    "As per policy document",
            "special_benefits":   "Refer to policy brochure",
            "raw_summary":        text[:200].replace("\n", " ") + "…",
        })

        log.info(f"[KB] Fallback extracted {len(plans)} plan(s)")
        return plans

    # ── MAIN UPLOAD PIPELINE ──────────────────────────────────────────────────
    def process_upload(self, filename, file_bytes, file_path, uploaded_by="admin") -> dict:
        """
        Full pipeline:
        1. Security scan
        2. Duplicate check
        3. Store document record
        4. Extract text
        5. AI extraction
        6. Store plans
        7. Mark active
        Returns: {"success": bool, "message": str, "doc_id": int, "plans": list}
        """
        ext = os.path.splitext(filename)[1].lower()

        # 1. Security scan
        ok, scan_msg = self._scan_file(file_bytes, ext)
        if not ok:
            return {"success": False, "message": scan_msg, "doc_id": None, "plans": []}

        if ext not in ALLOWED_EXTS:
            return {"success": False,
                    "message": f"Unsupported file type '{ext}'. Use: PDF, TXT, DOCX, JPG, PNG",
                    "doc_id": None, "plans": []}

        # 2. Duplicate check
        if self.db.kb_doc_exists(file_bytes):
            return {"success": False,
                    "message": "⚠️ This document already exists in the knowledge base.",
                    "doc_id": None, "plans": []}

        # 3. Store document record
        doc_id = self.db.kb_store_document(filename, file_path, file_bytes, uploaded_by)
        log.info(f"[KB] Stored doc_id={doc_id} '{filename}'")

        # 4. Extract text
        text = self._extract_text(file_bytes, ext, file_path)
        if not text:
            self.db.kb_update_doc_status(doc_id, 'failed_extraction')
            return {"success": False,
                    "message": "⚠️ Could not extract text from document. Try a text-based PDF or DOCX.",
                    "doc_id": doc_id, "plans": []}

        # 5. AI structured extraction
        plans = self._extract_plans_via_ai(text)
        if not plans:
            self.db.kb_update_doc_status(doc_id, 'failed_extraction')
            return {"success": False,
                    "message": "⚠️ Could not extract policy details. Ensure the document contains insurance plan information.",
                    "doc_id": doc_id, "plans": []}

        # 6. Store plans
        self.db.kb_store_plans(doc_id, plans, is_master=0)
        log.info(f"[KB] Stored {len(plans)} plans for doc_id={doc_id}")

        # 7. Mark active
        self.db.kb_update_doc_status(doc_id, 'active')

        return {
            "success":  True,
            "message":  f"✅ Document processed! Extracted {len(plans)} insurance plan(s) into knowledge base.",
            "doc_id":   doc_id,
            "plans":    plans,
        }

    # ── UPDATE PIPELINE ────────────────────────────────────────────────────────
    def process_update(self, doc_id, filename, file_bytes, file_path, note="") -> dict:
        """Re-process an existing document — keeps version history."""
        ext = os.path.splitext(filename)[1].lower()

        ok, scan_msg = self._scan_file(file_bytes, ext)
        if not ok:
            return {"success": False, "message": scan_msg}

        existing = self.db.kb_get_doc(doc_id)
        if not existing:
            return {"success": False, "message": "Document not found."}

        new_version = (existing.get("version") or 1) + 1
        self.db.kb_save_version(doc_id, existing["filename"], existing["file_path"],
                                new_version - 1, note or "Previous version")

        text = self._extract_text(file_bytes, ext, file_path)
        if not text:
            return {"success": False, "message": "⚠️ Could not extract text from updated document."}

        plans = self._extract_plans_via_ai(text)
        if not plans:
            return {"success": False, "message": "⚠️ Could not extract policy details from updated document."}

        self.db.kb_store_plans(doc_id, plans, is_master=0)
        self.db.kb_save_version(doc_id, filename, file_path, new_version, "Updated")
        self.db.kb_update_doc_status(doc_id, 'active')

        return {
            "success": True,
            "message": f"✅ Document updated! Re-extracted {len(plans)} plan(s). Version {new_version}.",
            "plans":   plans,
        }

    # ── RECOMMENDATION ENGINE ─────────────────────────────────────────────────
    def get_recommendations(self, user_profile: dict, top_n=3) -> list:
        """
        Filter and rank plans from DB based on user profile.
        Returns top_n best-matching plans with reasons.
        """
        all_plans = self.db.kb_get_all_plans_for_recommendation()
        if not all_plans:
            log.warning("[KB] No plans in database for recommendation")
            return []

        age       = user_profile.get("age", 0) or 0
        budget    = user_profile.get("budget_range", "") or ""
        ins_type  = user_profile.get("insurance_type", "") or ""
        medical   = (user_profile.get("medical_conditions", "") or "").lower()
        city      = (user_profile.get("city", "") or "").lower()
        family    = user_profile.get("family_members", "") or ""

        # Monthly budget parsing
        budget_max = 9999
        budget_map = {
            "under ₹500": 500, "₹500–₹1,000": 1000, "₹500-₹1,000": 1000,
            "₹1,000–₹2,000": 2000, "₹1,000-₹2,000": 2000,
            "₹2,000–₹5,000": 5000, "₹2,000-₹5,000": 5000,
            "above ₹5,000": 99999,
        }
        for k, v in budget_map.items():
            if k.lower() in budget.lower():
                budget_max = v
                break

        scored = []
        for plan in all_plans:
            score  = 0
            reason = []
            ptype  = (plan.get("insurance_type", "") or "").lower()
            pname  = (plan.get("plan_name", "") or "")
            pcov   = (plan.get("conditions_covered", "") or "").lower()
            pelig  = (plan.get("eligibility_age", "") or "").lower()
            pprem  = (plan.get("premium_range", "") or "").lower()
            phosp  = (plan.get("network_hospitals", "") or "").lower()
            pbene  = (plan.get("special_benefits", "") or "").lower()
            pwait  = (plan.get("waiting_period", "") or "").lower()

            # Insurance type match — STRICT filter (skip plans of wrong type)
            if ins_type:
                # Normalize for matching
                ins_norm = ins_type.lower().replace("term / life", "life").replace("term life", "life")
                type_match = (ins_norm in ptype) or (ptype in ins_norm)
                if not type_match:
                    continue   # Hard skip — never recommend wrong type
                score += 40
                reason.append(f"matches {ins_type}")
            else:
                score += 10

            # Medical condition matching
            if medical and medical not in ("none", "no", ""):
                for cond in ["diabetes", "hypertension", "bp", "heart", "asthma", "kidney", "cancer"]:
                    if cond in medical and cond in pcov:
                        score += 30
                        reason.append(f"covers {cond}")
                    if cond in medical and cond in pname.lower():
                        score += 20
                        reason.append(f"plan designed for {cond}")
                if "zero" in pwait or "day-1" in pwait or "immediate" in pwait:
                    score += 15
                    reason.append("immediate coverage for conditions")

            # Age eligibility
            age_ok = True
            age_match = re.search(r'(\d+)\s*[-–to]+\s*(\d+)', pelig)
            if age_match and age:
                mn, mx = int(age_match.group(1)), int(age_match.group(2))
                if mn <= int(age) <= mx:
                    score += 15
                    reason.append(f"eligible at age {age}")
                else:
                    score -= 20
                    age_ok = False

            # Senior citizen check
            if age and int(age) >= 60:
                if "senior" in pname.lower() or "red carpet" in pname.lower():
                    score += 25
                    reason.append("designed for senior citizens")

            # Family plan matching
            if family and ("spouse" in family.lower() or "child" in family.lower() or "parent" in family.lower()):
                if "family" in pname.lower() or "floater" in pname.lower():
                    score += 20
                    reason.append("family floater plan")

            # City/network hospital matching
            if city and city in phosp:
                score += 10
                reason.append(f"network hospitals in {city}")

            # Master document plans get slight boost
            if plan.get("is_master"):
                score += 5

            # Budget rough match (parse premium string)
            prem_nums = re.findall(r'\d[\d,]*', pprem.replace(",", ""))
            if prem_nums:
                try:
                    min_prem = int(prem_nums[0])
                    if min_prem <= budget_max:
                        score += 10
                        reason.append(f"fits ₹{budget_max} budget")
                    else:
                        score -= 5
                except Exception:
                    pass

            if score > 0 and age_ok:
                scored.append({
                    "plan":   plan,
                    "score":  score,
                    "reason": ", ".join(reason) if reason else "matches your profile",
                })

        # Sort by score descending
        scored.sort(key=lambda x: x["score"], reverse=True)
        top = scored[:top_n]

        # Log recommendation events and increment counters
        for item in top:
            pid = item["plan"].get("id")
            pname = item["plan"].get("plan_name", "")
            if pid:
                self.db.kb_increment_recommend(pid)
                self.db.kb_log_event("recommendation", pid, pname,
                                     user_profile.get("user_id",""), item["reason"])

        return top

    def format_recommendation_text(self, recommendations: list, user_profile: dict) -> str:
        """Format recommendations into a natural language response for the user."""
        if not recommendations:
            self.db.kb_log_event("failed_search", detail=
                f"ins={user_profile.get('insurance_type','')} age={user_profile.get('age','')} "
                f"budget={user_profile.get('budget_range','')}")
            return (
                "I wasn't able to find a matching plan in our knowledge base for your profile. "
                "Our advisor will contact you with personalised options! 📞"
            )

        name    = user_profile.get("name", "")
        age     = user_profile.get("age", "")
        budget  = user_profile.get("budget_range", "")
        medical = user_profile.get("medical_conditions", "")
        city    = user_profile.get("city", "")

        intro = f"Based on your profile"
        parts = []
        if name:  parts.append(f"{name}")
        if age:   parts.append(f"age {age}")
        if budget: parts.append(f"budget {budget}")
        if city:  parts.append(f"located in {city}")
        if medical and medical.lower() not in ("none","no",""):
            parts.append(f"with {medical}")
        if parts:
            intro += " (" + ", ".join(parts) + ")"
        intro += ", here are my top recommendations:\n\n"

        lines = []
        emojis = ["🥇", "🥈", "🥉"]
        for i, item in enumerate(recommendations):
            p      = item["plan"]
            reason = item["reason"]
            emoji  = emojis[i] if i < len(emojis) else "✅"

            line = (
                f"{emoji} **{p.get('plan_name','Plan')}** by {p.get('company_name','')}\n"
                f"   📋 {p.get('insurance_type','')} | "
                f"Coverage: {p.get('coverage_amount','N/A')} | "
                f"Premium: {p.get('premium_range','N/A')}\n"
                f"   ⏱️ Waiting: {p.get('waiting_period','N/A')} | "
                f"Age: {p.get('eligibility_age','N/A')}\n"
            )
            if p.get("special_benefits") and p["special_benefits"] != "Not specified":
                line += f"   ⭐ {p['special_benefits'][:100]}\n"
            line += f"   ✅ Why this plan: {reason}\n"
            lines.append(line)

        outro = "\nWould you like more details about any of these plans, or would you like to speak to an advisor? 😊"
        return intro + "\n".join(lines) + outro