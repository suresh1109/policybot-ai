"""
fix_kb_extraction.py
====================
Run this ONCE on your server to:
  1. Patch policy_kb.py (DOCX table extraction + better chunking + better prompt)
  2. Delete stale __pycache__ so Python uses the new code immediately
  3. Test the extraction on any .docx file you pass as argument

Usage:
  python fix_kb_extraction.py                          # just patches the files
  python fix_kb_extraction.py path/to/document.docx   # patches + tests extraction
"""

import os, sys, re, shutil, importlib

BASE = os.path.dirname(os.path.abspath(__file__))
KB_FILE   = os.path.join(BASE, "models", "policy_kb.py")
PYC_FILE  = os.path.join(BASE, "models", "__pycache__", "policy_kb.cpython-312.pyc")
PYC_FILE2 = os.path.join(BASE, "models", "__pycache__", "policy_kb.cpython-311.pyc")
PYC_FILE3 = os.path.join(BASE, "models", "__pycache__", "policy_kb.cpython-310.pyc")

# ── PATCH 1: DOCX extraction — read tables not just paragraphs ──────────────

OLD_DOCX = '''            elif ext in (".docx", ".doc") and DOCX_OK:
                import docx as docx_lib, io
                doc = docx_lib.Document(io.BytesIO(file_bytes))
                text = "\\n".join(p.text for p in doc.paragraphs if p.text.strip())
                log.info(f"[KB] DOCX: {len(text)} chars")'''

NEW_DOCX = '''            elif ext in (".docx", ".doc") and DOCX_OK:
                import docx as docx_lib, io
                from docx.oxml.ns import qn as _qn
                doc = docx_lib.Document(io.BytesIO(file_bytes))
                parts = []
                # Walk body in reading order so tables are included.
                # doc.paragraphs alone SILENTLY SKIPS ALL TABLE CONTENT —
                # every plan stored in a table becomes invisible to the AI.
                for child in doc.element.body:
                    tag = child.tag.split(\'}\')[-1] if \'}\' in child.tag else child.tag
                    if tag == \'p\':
                        para_text = \'\'.join(
                            r.text for r in child.iter(_qn(\'w:t\'))
                        ).strip()
                        if para_text:
                            parts.append(para_text)
                    elif tag == \'tbl\':
                        rows = child.findall(\'.//' + _qn(\'w:tr\') + \'\')
                        for ri, row in enumerate(rows):
                            cells_text = []
                            for cell in row.findall(\'.//' + _qn(\'w:tc\') + \'\'):
                                cell_val = \' \'.join(
                                    \'\'.join(r2.text for r2 in p2.iter(_qn(\'w:t\'))).strip()
                                    for p2 in cell.findall(\'.//' + _qn(\'w:p\') + \'\')
                                ).strip()
                                if cell_val:
                                    cells_text.append(cell_val)
                            if cells_text:
                                prefix = \'HEADER: \' if ri == 0 else \'\'
                                parts.append(prefix + \' | \'.join(cells_text))
                text = \'\\n\'.join(parts)
                log.info(f"[KB] DOCX: {len(text)} chars ({len(doc.paragraphs)} paras + {len(doc.tables)} tables)")'''

# ── PATCH 2: Chunk size and output tokens ────────────────────────────────────

OLD_CHUNK = "        MAX_CHUNK = 14000"
NEW_CHUNK = "        MAX_CHUNK  = 18000\n        MAX_TOKENS = 7000"

OLD_STEP = "            step = MAX_CHUNK - 1000"
NEW_STEP = "            step = MAX_CHUNK - 500"

OLD_TOKENS = "                response = self.gemini.generate(prompt, max_tokens=4096)"
NEW_TOKENS = "                response = self.gemini.generate(prompt, max_tokens=MAX_TOKENS)"

# ── PATCH 3: Better extraction prompt ───────────────────────────────────────

OLD_PROMPT = '''EXTRACTION_PROMPT = """You are an expert Indian insurance policy data extractor.
Read the policy document text below and extract structured plan information.

Return ONLY a valid JSON array. Each element is one insurance plan with EXACTLY these keys:
  company_name, plan_name, insurance_type, coverage_amount, premium_range,
  waiting_period, conditions_covered, exclusions, claim_process,
  network_hospitals, eligibility_age, special_benefits, raw_summary

EXTRACTION RULES:
- insurance_type: MUST be one of exactly: Health Insurance, Life Insurance, Term Life Insurance,
  Vehicle Insurance, Travel Insurance, Property Insurance, Accident Insurance
- company_name: Look for insurer name (Star Health, HDFC ERGO, Bajaj Allianz, ICICI Lombard,
  LIC, SBI Life, Tata AIG, Niva Bupa, Care Health, Aditya Birla, New India, United India, etc.)
- plan_name: Find the exact product/plan name (e.g., "Optima Restore", "Click2Protect", "Comprehensive Motor")
- coverage_amount: Sum insured or coverage limit (e.g., "₹5 Lakh to ₹1 Crore")
- premium_range: Monthly or annual premium (e.g., "₹800–₹2,500/month")
- waiting_period: Initial waiting period (e.g., "30 days", "2 years for pre-existing")
- conditions_covered: Diseases/conditions covered (comma-separated)
- exclusions: What is NOT covered
- claim_process: How to file a claim (cashless/reimbursement, helpline, etc.)
- network_hospitals: Hospital network size or names (e.g., "10,000+ hospitals")
- eligibility_age: Age eligibility (e.g., "18–65 years")
- special_benefits: Unique benefits (no-claim bonus, free health check, etc.)
- raw_summary: 2-3 sentence plain English summary of this specific plan

IMPORTANT:
- If a field is not found in the document, write exactly: "Not specified"
- Extract ALL distinct plans/products found — a document may have multiple plans
- For vehicle insurance: coverage_amount = IDV (Insured Declared Value), premium varies by vehicle
- Return ONLY the JSON array, no explanation, no markdown, no code fences

Document text:
{text}
"""'''

NEW_PROMPT = '''EXTRACTION_PROMPT = """You are an expert Indian insurance policy data extractor.
Read the document text below and extract EVERY insurance plan found.

This document is a CATALOG — it may contain 60 or more plans across many sections.
You MUST extract EVERY SINGLE PLAN. Do not stop early.

Return ONLY a valid JSON array. Each element is one insurance plan with EXACTLY these keys:
  company_name, plan_name, insurance_type, coverage_amount, premium_range,
  waiting_period, conditions_covered, exclusions, claim_process,
  network_hospitals, eligibility_age, special_benefits, raw_summary

EXTRACTION RULES:
- insurance_type: MUST be one of: Health Insurance, Life Insurance, Term Life Insurance,
  Vehicle Insurance, Travel Insurance, Property Insurance, Accident Insurance
- company_name: The insurer name (Star Health, HDFC ERGO, Bajaj Allianz, ICICI Lombard,
  LIC India, SBI Life, Tata AIG, Niva Bupa, Care Health, Aditya Birla, New India,
  United India, Oriental Insurance, Reliance General, Digit Insurance, ICICI Prudential,
  Max Life, Kotak, Royal Sundaram, Manipal Cigna, National Insurance, All Insurers, etc.)
- plan_name: The exact product name as written
- coverage_amount: Sum insured / coverage limit (e.g. "Rs.5 Lakh to Rs.1 Crore")
- premium_range: Annual or monthly premium (e.g. "Rs.8,000-Rs.25,000/year")
- waiting_period: PED or general waiting period; "Not specified" if absent
- conditions_covered: Diseases/events covered comma-separated
- exclusions: What is NOT covered; "Not specified" if absent
- claim_process: How to claim; "Not specified" if absent
- network_hospitals: Network size or "Not specified"
- eligibility_age: Age range or "Not specified"
- special_benefits: Unique features or "Not specified"
- raw_summary: 1-2 sentence plain English summary of this specific plan

HOW TO READ TABLE ROWS:
- Lines starting with HEADER: list the column names for the rows that follow
- All other pipe-separated lines are data rows: val1 | val2 | val3 ...
- Map each value to the column name above it
- EACH data row in a plan table = ONE separate plan object in your output

EXAMPLE:
  HEADER: Plan Name | Insurer | Sum Insured | Annual Premium | Key Feature
  Arogya Sanjeevani | All Insurers | Rs.1L-Rs.5L | Rs.3,000-Rs.8,000 | Standard policy
  MyHealth Suraksha | HDFC ERGO | Rs.3L-Rs.75L | Rs.5,500-Rs.22,000 | No room rent cap
→ Output TWO plan objects: one for Arogya Sanjeevani and one for MyHealth Suraksha.

IMPORTANT:
- If a field is not found, write exactly: "Not specified"
- Do NOT merge multiple plans into one entry
- Return ONLY the JSON array — no preamble, no explanation, no markdown fences

Document text:
{text}
"""'''


def apply_patches():
    print(f"\n[1] Reading {KB_FILE}")
    with open(KB_FILE, "r", encoding="utf-8") as f:
        src = f.read()

    original = src
    applied = []

    # Apply each patch
    patches = [
        ("DOCX table extraction", OLD_DOCX, NEW_DOCX),
        ("Extraction prompt",     OLD_PROMPT, NEW_PROMPT),
        ("MAX_CHUNK size",        OLD_CHUNK, NEW_CHUNK),
        ("Overlap step",          OLD_STEP,  NEW_STEP),
        ("Output tokens",         OLD_TOKENS, NEW_TOKENS),
    ]

    for name, old, new in patches:
        if old in src:
            src = src.replace(old, new)
            applied.append(f"  ✅  {name}")
        elif new in src:
            applied.append(f"  ✔  {name} (already applied)")
        else:
            applied.append(f"  ⚠  {name} — pattern not found, skipping")

    for msg in applied:
        print(msg)

    if src != original:
        # Backup original
        backup = KB_FILE + ".backup"
        shutil.copy2(KB_FILE, backup)
        print(f"\n[2] Backup saved: {backup}")

        with open(KB_FILE, "w", encoding="utf-8") as f:
            f.write(src)
        print(f"[3] policy_kb.py updated ✅")
    else:
        print("\n[2] No changes needed (all patches already applied)")

    # Delete ALL pycache versions
    print("\n[4] Clearing __pycache__...")
    cleared = 0
    pycache_dir = os.path.join(BASE, "models", "__pycache__")
    if os.path.isdir(pycache_dir):
        for fname in os.listdir(pycache_dir):
            if "policy_kb" in fname:
                fpath = os.path.join(pycache_dir, fname)
                os.remove(fpath)
                print(f"  🗑  Deleted: {fname}")
                cleared += 1
    if cleared == 0:
        print("  (no policy_kb pycache files found)")
    else:
        print(f"  ✅  Cleared {cleared} cached bytecode file(s)")

    print("\n✅ Patch complete.")
    print("👉 NOW RESTART FLASK:  Ctrl+C  then  python app.py")
    print("   Then click Re-extract (✨) on the document in Admin → Policy KB\n")


def test_extraction(docx_path):
    """Run the new extraction on a docx file and show what Gemini will receive."""
    print(f"\n[TEST] Testing extraction on: {docx_path}")
    try:
        import docx as docx_lib, io
        from docx.oxml.ns import qn as _qn

        with open(docx_path, "rb") as f:
            file_bytes = f.read()

        doc = docx_lib.Document(io.BytesIO(file_bytes))
        parts = []
        for child in doc.element.body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'p':
                para_text = ''.join(r.text for r in child.iter(_qn('w:t'))).strip()
                if para_text:
                    parts.append(para_text)
            elif tag == 'tbl':
                rows = child.findall('.//' + _qn('w:tr'))
                for ri, row in enumerate(rows):
                    cells_text = []
                    for cell in row.findall('.//' + _qn('w:tc')):
                        cell_val = ' '.join(
                            ''.join(r2.text for r2 in p2.iter(_qn('w:t'))).strip()
                            for p2 in cell.findall('.//' + _qn('w:p'))
                        ).strip()
                        if cell_val:
                            cells_text.append(cell_val)
                    if cells_text:
                        parts.append(('HEADER: ' if ri == 0 else '') + ' | '.join(cells_text))

        text = '\n'.join(parts)
        plan_rows = [l for l in text.split('\n')
                     if ' | ' in l and not l.startswith('HEADER:')
                     and len(l.split(' | ')) >= 3]

        print(f"  Total text extracted  : {len(text):,} chars")
        print(f"  Paragraphs found      : {len(doc.paragraphs)}")
        print(f"  Tables found          : {len(doc.tables)}")
        print(f"  Plan data rows visible: {len(plan_rows)}")
        print(f"  Chunks for Gemini     : {max(1, (len(text)-1) // 17500 + 1)}")
        print()
        print("  First 10 plan rows:")
        for r in plan_rows[:10]:
            print(f"    {r[:100]}")
        if len(plan_rows) > 10:
            print(f"    ... and {len(plan_rows)-10} more")
        print()
        if len(plan_rows) > 0:
            print(f"  ✅ Extraction will work — {len(plan_rows)} rows will be sent to Gemini")
        else:
            print("  ⚠  No pipe-delimited rows found — document may not have tables")
    except Exception as e:
        print(f"  ❌ Test error: {e}")
        import traceback; traceback.print_exc()


if __name__ == "__main__":
    apply_patches()

    # If a docx path was passed as argument, test it
    if len(sys.argv) > 1:
        test_extraction(sys.argv[1])